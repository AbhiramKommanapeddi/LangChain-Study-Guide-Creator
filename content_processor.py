"""
Content processor for extracting and analyzing text from various document formats.
Handles PDF parsing, text cleaning, and concept identification.
"""

import os
import re
import nltk
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path

import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
import numpy as np

try:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
    nltk.download('averaged_perceptron_tagger', quiet=True)
except:
    pass

@dataclass
class ProcessedContent:
    """Container for processed document content."""
    text: str
    chunks: List[str]
    concepts: List[str]
    key_terms: List[str]
    metadata: Dict
    sections: List[Dict]

class ContentProcessor:
    """Processes various document formats and extracts educational content."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
    def process_document(self, file_path: str, document_type: str = "auto") -> ProcessedContent:
        """
        Process a document and extract educational content.
        
        Args:
            file_path: Path to the document
            document_type: Type of document (pdf, docx, txt, auto)
            
        Returns:
            ProcessedContent object with extracted information
        """
        if document_type == "auto":
            document_type = self._detect_file_type(file_path)
            
        # Extract raw text
        if document_type == "pdf":
            text = self._extract_from_pdf(file_path)
        elif document_type == "docx":
            text = self._extract_from_docx(file_path)
        elif document_type == "txt":
            text = self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported document type: {document_type}")
            
        # Clean and process text
        cleaned_text = self._clean_text(text)
        
        # Split into chunks
        chunks = self.text_splitter.split_text(cleaned_text)
        
        # Extract concepts and key terms
        concepts = self._extract_concepts(cleaned_text)
        key_terms = self._extract_key_terms(cleaned_text)
        
        # Identify sections
        sections = self._identify_sections(cleaned_text)
        
        # Create metadata
        metadata = {
            "file_path": file_path,
            "file_type": document_type,
            "word_count": len(cleaned_text.split()),
            "chunk_count": len(chunks),
            "concept_count": len(concepts)
        }
        
        return ProcessedContent(
            text=cleaned_text,
            chunks=chunks,
            concepts=concepts,
            key_terms=key_terms,
            metadata=metadata,
            sections=sections
        )
    
    def _detect_file_type(self, file_path: str) -> str:
        """Detect file type from extension."""
        extension = Path(file_path).suffix.lower()
        if extension == ".pdf":
            return "pdf"
        elif extension in [".docx", ".doc"]:
            return "docx"
        elif extension == ".txt":
            return "txt"
        else:
            raise ValueError(f"Unsupported file extension: {extension}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            raise ValueError(f"Error reading PDF file: {str(e)}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise ValueError(f"Error reading DOCX file: {str(e)}")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise ValueError(f"Error reading TXT file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)]', ' ', text)
        
        # Fix common OCR errors
        text = re.sub(r'\s+([\.!?])', r'\1', text)
        text = re.sub(r'([\.!?])\s*([A-Z])', r'\1 \2', text)
        
        return text.strip()
    
    def _extract_concepts(self, text: str) -> List[str]:
        """Extract key concepts using TF-IDF and clustering."""
        try:
            # Split into sentences
            sentences = nltk.sent_tokenize(text)
            
            if len(sentences) < 5:
                return []
            
            # Use TF-IDF to find important terms
            vectorizer = TfidfVectorizer(
                max_features=100,
                stop_words='english',
                ngram_range=(1, 3),
                min_df=2
            )
            
            tfidf_matrix = vectorizer.fit_transform(sentences)
            feature_names = vectorizer.get_feature_names_out()
            
            # Get top terms by TF-IDF score
            scores = np.array(tfidf_matrix.sum(axis=0)).flatten()
            top_indices = scores.argsort()[-20:][::-1]
            
            concepts = [feature_names[i] for i in top_indices if scores[i] > 0.1]
            
            return concepts[:10]  # Return top 10 concepts
            
        except Exception as e:
            print(f"Error extracting concepts: {e}")
            return []
    
    def _extract_key_terms(self, text: str) -> List[str]:
        """Extract key terms using POS tagging."""
        try:
            # Tokenize and tag
            tokens = nltk.word_tokenize(text)
            pos_tags = nltk.pos_tag(tokens)
            
            # Extract nouns and noun phrases
            key_terms = []
            for word, pos in pos_tags:
                if pos in ['NN', 'NNP', 'NNS', 'NNPS'] and len(word) > 3:
                    key_terms.append(word.lower())
            
            # Count frequency and return most common
            from collections import Counter
            term_counts = Counter(key_terms)
            
            return [term for term, count in term_counts.most_common(20)]
            
        except Exception as e:
            print(f"Error extracting key terms: {e}")
            return []
    
    def _identify_sections(self, text: str) -> List[Dict]:
        """Identify document sections and chapters."""
        sections = []
        
        # Look for common section patterns
        section_patterns = [
            r'^Chapter\s+\d+',
            r'^\d+\.\s+',
            r'^[A-Z][A-Z\s]+$',
            r'^[A-Z][a-z]+:',
        ]
        
        lines = text.split('\n')
        current_section = None
        section_content = []
        
        for line_num, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # Check if this line is a section header
            is_header = False
            for pattern in section_patterns:
                if re.match(pattern, line):
                    is_header = True
                    break
            
            if is_header:
                # Save previous section
                if current_section:
                    sections.append({
                        'title': current_section,
                        'content': '\n'.join(section_content),
                        'start_line': current_section_start,
                        'end_line': line_num - 1
                    })
                
                # Start new section
                current_section = line
                current_section_start = line_num
                section_content = []
            else:
                section_content.append(line)
        
        # Add final section
        if current_section:
            sections.append({
                'title': current_section,
                'content': '\n'.join(section_content),
                'start_line': current_section_start,
                'end_line': len(lines) - 1
            })
        
        return sections

    def create_langchain_documents(self, processed_content: ProcessedContent) -> List[LangChainDocument]:
        """Convert processed content to LangChain documents."""
        documents = []
        
        for i, chunk in enumerate(processed_content.chunks):
            metadata = {
                **processed_content.metadata,
                'chunk_id': i,
                'concepts': processed_content.concepts[:5],  # Include top concepts
                'key_terms': processed_content.key_terms[:10]  # Include top terms
            }
            
            doc = LangChainDocument(
                page_content=chunk,
                metadata=metadata
            )
            documents.append(doc)
        
        return documents
